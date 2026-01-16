"""DOCX text extraction."""

from typing import List, Dict
import logging

logger = logging.getLogger(__name__)


def extract_text_from_docx(file_path: str) -> List[Dict]:
    """Extract text from DOCX file."""
    try:
        from docx import Document as DocxDocument
        
        doc = DocxDocument(file_path)
        pages = []
        
        # DOCX doesn't have pages, so we'll treat paragraphs as "pages"
        # Group paragraphs into logical pages (every ~500 words)
        current_page_text = []
        current_page_num = 1
        word_count = 0
        words_per_page = 500
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                words = text.split()
                word_count += len(words)
                current_page_text.append(text)
                
                # If we've accumulated enough words, create a page
                if word_count >= words_per_page:
                    pages.append({
                        "page_number": current_page_num,
                        "text": "\n\n".join(current_page_text)
                    })
                    current_page_text = []
                    current_page_num += 1
                    word_count = 0
        
        # Add remaining text as last page
        if current_page_text:
            pages.append({
                "page_number": current_page_num,
                "text": "\n\n".join(current_page_text)
            })
        
        # If no pages were created, create one with all text
        if not pages:
            all_text = "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            if all_text:
                pages.append({
                    "page_number": 1,
                    "text": all_text
                })
        
        return pages if pages else [{"page_number": 1, "text": ""}]
        
    except ImportError:
        logger.error("python-docx not installed. Install it with: pip install python-docx")
        raise Exception("DOCX processing requires python-docx package")
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")

