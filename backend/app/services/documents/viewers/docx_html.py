"""DOCX to HTML conversion service for viewing."""

import logging
from typing import Optional

logger = logging.getLogger(__name__)


def convert_docx_to_html(file_path: str) -> str:
    """
    Convert DOCX file to HTML for viewing.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        HTML content as string
        
    Raises:
        Exception: If conversion fails
    """
    try:
        from docx import Document as DocxDocument
        
        doc = DocxDocument(file_path)
        
        html_parts = []
        html_parts.append('<!DOCTYPE html>')
        html_parts.append('<html><head>')
        html_parts.append('<meta charset="UTF-8">')
        html_parts.append('<meta name="viewport" content="width=device-width, initial-scale=1.0">')
        html_parts.append('<style>')
        html_parts.append('''
            body {
                font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif;
                max-width: 900px;
                margin: 0 auto;
                padding: 2rem;
                line-height: 1.6;
                color: #333;
                background: #f5f5f5;
            }
            .document {
                background: white;
                padding: 3rem;
                box-shadow: 0 2px 8px rgba(0,0,0,0.1);
                border-radius: 8px;
            }
            p {
                margin: 0.75rem 0;
            }
            h1, h2, h3, h4, h5, h6 {
                margin-top: 1.5rem;
                margin-bottom: 0.75rem;
                font-weight: 600;
            }
            h1 { font-size: 2rem; }
            h2 { font-size: 1.75rem; }
            h3 { font-size: 1.5rem; }
            h4 { font-size: 1.25rem; }
            h5 { font-size: 1.1rem; }
            h6 { font-size: 1rem; }
            ul, ol {
                margin: 0.75rem 0;
                padding-left: 2rem;
            }
            li {
                margin: 0.25rem 0;
            }
            table {
                border-collapse: collapse;
                width: 100%;
                margin: 1rem 0;
            }
            table td, table th {
                border: 1px solid #ddd;
                padding: 0.5rem;
                text-align: left;
            }
            table th {
                background-color: #f2f2f2;
                font-weight: 600;
            }
            .text-center { text-align: center; }
            .text-right { text-align: right; }
            .text-left { text-align: left; }
            strong { font-weight: 600; }
            em { font-style: italic; }
            u { text-decoration: underline; }
        ''')
        html_parts.append('</style>')
        html_parts.append('</head><body>')
        html_parts.append('<div class="document">')
        
        # Process all paragraphs
        for paragraph in doc.paragraphs:
            html_parts.append(process_paragraph_simple(paragraph))
        
        # Process all tables
        for table in doc.tables:
            html_parts.append(process_table_simple(table))
        
        html_parts.append('</div>')
        html_parts.append('</body></html>')
        
        return '\n'.join(html_parts)
        
    except ImportError:
        logger.error("python-docx not installed")
        raise Exception("DOCX processing requires python-docx package")
    except Exception as e:
        logger.error(f"Error converting DOCX to HTML {file_path}: {str(e)}")
        raise Exception(f"Failed to convert DOCX to HTML: {str(e)}")


def process_paragraph_simple(paragraph):
    """Process a paragraph and return HTML."""
    if not paragraph.text.strip() and len(paragraph.runs) == 0:
        return '<p>&nbsp;</p>'
    
    # Determine tag based on style
    style_name = paragraph.style.name if paragraph.style else ''
    tag = 'p'
    if 'Heading 1' in style_name or 'heading 1' in style_name.lower():
        tag = 'h1'
    elif 'Heading 2' in style_name or 'heading 2' in style_name.lower():
        tag = 'h2'
    elif 'Heading 3' in style_name or 'heading 3' in style_name.lower():
        tag = 'h3'
    elif 'Heading 4' in style_name or 'heading 4' in style_name.lower():
        tag = 'h4'
    elif 'Heading 5' in style_name or 'heading 5' in style_name.lower():
        tag = 'h5'
    elif 'Heading 6' in style_name or 'heading 6' in style_name.lower():
        tag = 'h6'
    
    # Check alignment
    align_class = ''
    if paragraph.alignment:
        if paragraph.alignment == 1:  # CENTER
            align_class = ' class="text-center"'
        elif paragraph.alignment == 2:  # RIGHT
            align_class = ' class="text-right"'
    
    # Process runs (text with formatting)
    content_parts = []
    for run in paragraph.runs:
        if run.text:
            formatted_text = run.text
            if run.bold:
                formatted_text = f'<strong>{formatted_text}</strong>'
            if run.italic:
                formatted_text = f'<em>{formatted_text}</em>'
            if run.underline:
                formatted_text = f'<u>{formatted_text}</u>'
            content_parts.append(formatted_text)
    
    content = ''.join(content_parts) if content_parts else paragraph.text
    if not content.strip():
        content = '&nbsp;'
    
    return f'<{tag}{align_class}>{content}</{tag}>'


def process_table_simple(table):
    """Process a table and return HTML."""
    html_parts = ['<table>']
    
    for row_idx, row in enumerate(table.rows):
        html_parts.append('<tr>')
        tag = 'th' if row_idx == 0 else 'td'
        
        for cell in row.cells:
            cell_text = []
            for paragraph in cell.paragraphs:
                if paragraph.text.strip():
                    cell_text.append(paragraph.text)
            
            cell_content = ' '.join(cell_text) if cell_text else '&nbsp;'
            html_parts.append(f'<{tag}>{cell_content}</{tag}>')
        
        html_parts.append('</tr>')
    
    html_parts.append('</table>')
    return ''.join(html_parts)

